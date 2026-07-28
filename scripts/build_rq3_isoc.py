"""
Model 1b: Chinese CDN market concentration (CDN_China) built from the ISOC dataset --
the companion to the Nedko-based Model 1. Same specification: FE + RE, cluster-robust SEs
by country, Hausman test, wild-cluster-bootstrap robustness.

CDN_China (ISOC) = share (%) of top-list websites served by a Chinese-jurisdiction CDN
(in ISOC that is 'alibaba'), website-count-weighted (local+external, latest date per month),
mean of the monthly shares per country-year. Years 2024-2026, IVs lagged one year.

Writes to rq3_regression_isoc.docx -- a SEPARATE file. It does NOT open or overwrite
rq3_regression.docx.

Run: /opt/anaconda3/bin/python3 build_rq3_isoc.py
"""
import csv, os
from collections import defaultdict
import numpy as np
from docx import Document
from docx.shared import Pt
from build_rq3_regression import (run_model, wcb_p, hinges_on, fstat_str, stars,
                                  XCOLS, VARROWS, CCS)

BASE = "/Users/newlivehung/Desktop/11. Pulse Research Fellowship/06.07.2026_Folder"
RAW_ISOC = "/Users/newlivehung/Desktop/11. Pulse Research Fellowship/26.05.2026_Folder/lccrawresults.csv"
SEA = set(CCS)
DROP = {"summary", "other"}
CHINA_ISOC = {"alibaba"}          # the only Chinese-jurisdiction host_cdn in ISOC
ISOC_WEIGHT = "local+external"     # same weight used in the concentration analysis


def isoc_cdn_china():
    """(cc,year) -> mean monthly China CDN share (%), ISOC website-count weighted."""
    tmp = {}   # (cc,month,host_cdn) -> (date_key, weight) keeping latest date in month
    with open(RAW_ISOC, newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            cc = (row.get("test_cc") or "").strip(); cdn = (row.get("host_cdn") or "").strip()
            if cc not in SEA or cdn in DROP or not cdn:
                continue
            dd, mm, yyyy = row["date"].split("/")
            month = yyyy + mm.zfill(2); dkey = (yyyy, mm.zfill(2), dd.zfill(2))
            l = int(row["local_websites"] or 0); e = int(row["external_websites"] or 0)
            w = l + e if ISOC_WEIGHT == "local+external" else (l if ISOC_WEIGHT == "local" else e)
            k = (cc, month, cdn)
            if k not in tmp or dkey > tmp[k][0]:
                tmp[k] = (dkey, w)
    china = defaultdict(float); total = defaultdict(float)
    for (cc, month, cdn), (dk, w) in tmp.items():
        total[(cc, month)] += w
        if cdn in CHINA_ISOC:
            china[(cc, month)] += w
    yr = defaultdict(list)
    for k, tot in total.items():
        if tot > 0:
            yr[(k[0], k[1][:4])].append(100 * china.get(k, 0.0) / tot)
    return {k: float(np.mean(v)) for k, v in yr.items()}


def plain_table(doc, res):
    fe, re = res["fe"], res["re"]
    t = doc.add_table(rows=1, cols=5); t.style = "Table Grid"
    for i, x in enumerate(["Variable", "FE coef", "FE SE", "RE coef", "RE SE"]):
        t.rows[0].cells[i].text = x
    for name, disp in VARROWS:
        c = t.add_row().cells; c[0].text = disp
        if name in fe.params.index:
            c[1].text = f"{fe.params[name]:.4f}{stars(fe.pvalues[name])}"; c[2].text = f"{fe.std_errors[name]:.4f}"
        else:
            c[1].text = "(absorbed)"
        c[3].text = f"{re.params[name]:.4f}{stars(re.pvalues[name])}"; c[4].text = f"{re.std_errors[name]:.4f}"
    c = t.add_row().cells; c[0].text = "constant"; c[1].text = "(absorbed)"
    c[3].text = f"{re.params['const']:.4f}{stars(re.pvalues['const'])}"; c[4].text = f"{re.std_errors['const']:.4f}"
    hs, hk, hp = res["hausman"]
    haus = f"chi2({hk})={hs:.2f}, p={hp:.3f}" if hs is not None else "degenerate (short panel)"
    for lbl, a, b in [("N", str(res["N"]), str(res["N"])),
                      ("Groups", str(res["G"]), str(res["G"])),
                      ("R2 within", f"{fe.rsquared_within:.4f}", f"{re.rsquared_within:.4f}"),
                      ("R2 overall", f"{fe.rsquared_overall:.4f}", f"{re.rsquared_overall:.4f}"),
                      ("F / Wald", fstat_str(fe), fstat_str(re))]:
        c = t.add_row().cells; c[0].text = lbl; c[1].text = a; c[3].text = b
    doc.add_paragraph(f"Hausman test (FE vs RE): {haus}. Panel: {res['G']} countries × T={res['T']}, "
                      f"N={res['N']}. Cluster-robust SEs by country. "
                      f"Significance: *** p<0.001, ** p<0.01, * p<0.05.")


def plain_robust(doc, res):
    sig = [v for v in XCOLS if v in res["fe"].pvalues.index and res["fe"].pvalues[v] < 0.05]
    if not sig:
        doc.add_paragraph("No coefficient significant at p<0.05.")
        return
    t = doc.add_table(rows=1, cols=1 + len(sig)); t.style = "Table Grid"
    for i, v in enumerate(sig):
        t.rows[0].cells[i + 1].text = v
    pdf = res["pdf"]
    def r(lbl, vals):
        c = t.add_row().cells; c[0].text = lbl
        for i, x in enumerate(vals):
            c[i + 1].text = x
    wcb = [wcb_p(pdf, XCOLS, XCOLS.index(v)) for v in sig]
    r("naive p", [f"{res['fe'].pvalues[v]:.4f}" for v in sig])
    r("wild bootstrap p", [f"{p:.3f}" for p in wcb])
    r("p<0.05?", ["Yes" if p < 0.05 else "No" for p in wcb])
    r("hinges on", [hinges_on(pdf, XCOLS, XCOLS.index(v)) for v in sig])


def main():
    print("computing ISOC CDN_China and running Model 1b...")
    dv = isoc_cdn_china()
    res = run_model(dv, [2024, 2025, 2026], "Model 1b (CDN, ISOC)")
    doc = Document()
    doc.add_heading("Model 1b - Chinese CDN market concentration (CDN_China), ISOC source", 1)
    plain_table(doc, res)
    doc.add_heading("Robustness", 2)
    plain_robust(doc, res)
    out = os.path.join(BASE, "rq3_regression_isoc.docx")
    doc.save(out)
    # console summary
    fe = res["fe"]
    sig = [f"{v} p={fe.pvalues[v]:.3f}" for v in XCOLS if v in fe.pvalues.index and fe.pvalues[v] < 0.05]
    print(f"  N={res['N']}, groups={res['G']} | FE significant (p<0.05): {sig if sig else 'NONE'}")
    print(f"DONE -> {out}   (rq3_regression.docx NOT touched)")


if __name__ == "__main__":
    main()
