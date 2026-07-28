"""
RQ3 panel regression (you run this): does lagged Chinese DSR (CGIT) predict Chinese
CDN / IP-infrastructure market concentration across the 10 SEA countries?

Model 1: CDN_China_it = b0 + b1 CGIT_projects_{i,t-1} + b2 CGIT_amount_{i,t-1} + b3 CGIT_type_{i,t-1} + a_i + e_it
Model 2: IP_China_it  = (same RHS)

DV = China-jurisdiction market share (%) per country-year, mean of the monthly shares:
     Model 1  CDN_China : Nedko (China CDN share of served domains), years 2024-2026.
     Model 2  IP_China  : OpenINTEL (China share of distinct websites), years 2025-2026
                          (2025 = mean of the Mar-Dec 2025 monthly shares, per the draft).
IVs (lagged t-1) from AEI CGIT (technology sector), aggregated per country-year:
     CGIT_projects = number of tech projects; CGIT_amount = US$m; CGIT_type = 1 if ANY
     construction deal that year else 0 (investment).

Estimators: fixed effects (within) and random effects, cluster-robust SEs by country,
plus a Hausman test to choose between them. Output: rq3_regression.docx (plain tables).

IMPORTANT: short panels (T=3 Model 1, T=2 Model 2; ~5-6 DSR-active countries). Treat the
estimates as EXPLORATORY -- and treat any p<0.05 with caution (small N + sparse predictors
can yield spurious significance as easily as nulls). The point of running it is to see the
ACTUAL numbers, not to assume them.

Requires:  pip install linearmodels
Run:       /opt/anaconda3/bin/python3 build_rq3_regression.py
"""
import csv, os, sys
from collections import defaultdict
import numpy as np
import pandas as pd
from scipy import stats
from linearmodels.panel import PanelOLS, RandomEffects
from docx import Document
from docx.shared import Pt

csv.field_size_limit(sys.maxsize)
BASE = "/Users/newlivehung/Desktop/11. Pulse Research Fellowship/06.07.2026_Folder"
CCS = ["BN", "ID", "KH", "LA", "MM", "MY", "PH", "SG", "TH", "VN"]
NAME = {"BN": "Brunei", "ID": "Indonesia", "KH": "Cambodia", "LA": "Laos", "MM": "Myanmar",
        "MY": "Malaysia", "PH": "Philippines", "SG": "Singapore", "TH": "Thailand", "VN": "Vietnam"}


def to_china(v):
    return (v or "").strip() in ("China", "CN")


def nedko_cdn_china():
    """(cc,year) -> mean monthly China CDN share (%) among served domains."""
    monthly = defaultdict(lambda: [0, 0])  # (cc,ym) -> [china, total]
    for cc in CCS:
        p = os.path.join(BASE, "nedko_sea_output_mapped", f"nedko_{cc.lower()}_202401_202605.csv")
        for row in csv.DictReader(open(p, encoding="utf-8")):
            if row["cdn_provider"] in ("unknown", "None detected"):
                continue
            k = (cc, row["measurement_period"]); monthly[k][1] += 1
            if to_china(row["cdn_jurisdiction"]): monthly[k][0] += 1
    yr = defaultdict(list)
    for (cc, ym), (ch, tot) in monthly.items():
        if tot: yr[(cc, ym[:4])].append(100 * ch / tot)
    return {k: float(np.mean(v)) for k, v in yr.items()}


def openintel_ip_china():
    """(cc,year) -> mean monthly China IP share (%) of distinct served websites."""
    chw = defaultdict(set); allw = defaultdict(set)
    for cc in CCS:
        p = os.path.join(BASE, "OpenINTEL_sea_output_caida_mapped", f"openintel_ip_infrastructure_extended_{cc}.csv")
        with open(p, newline="", encoding="utf-8") as f:
            r = csv.reader(f); h = next(r)
            iy, iw, ij, ip = h.index("measurement_period"), h.index("website"), h.index("infrastructure_jurisdiction"), h.index("ip_infrastructure_provider")
            for row in r:
                if row[ip].startswith("ASN "): continue
                k = (cc, row[iy]); allw[k].add(row[iw])
                if to_china(row[ij]): chw[k].add(row[iw])
    yr = defaultdict(list)
    for k, s in allw.items():
        yr[(k[0], k[1][:4])].append(100 * len(chw.get(k, set())) / len(s))
    return {k: float(np.mean(v)) for k, v in yr.items()}


def cgit_by_country_year():
    """(cc_name-> we key by 2-letter) (cc,year) -> (projects, amount, type_dummy)."""
    agg = defaultdict(lambda: [0, 0.0, 0])  # projects, amount, type(any construction)
    name2cc = {v: k for k, v in NAME.items()}
    for row in csv.DictReader(open(os.path.join(BASE, "AEI.sea.output.csv"), encoding="utf-8")):
        cc = name2cc.get(row["country"].strip())
        if cc is None: continue
        proj = (row["projects"] or "").strip(); amt = (row["amount"] or "").strip()
        if not proj or proj == "0": continue
        key = (cc, row["year"].strip())
        agg[key][0] += int(proj)
        agg[key][1] += float(amt) if amt else 0.0
        if row["type"].strip() == "construction": agg[key][2] = 1
    return {k: (v[0], v[1], v[2]) for k, v in agg.items()}


def build_panel(dv, dv_years):
    cg = cgit_by_country_year()
    rows = []
    for cc in CCS:
        for y in dv_years:
            d = dv.get((cc, str(y)))
            if d is None: continue
            proj, amt, typ = cg.get((cc, str(y - 1)), (0, 0.0, 0))
            rows.append((cc, y, d, proj, amt, typ))
    return pd.DataFrame(rows, columns=["country", "year", "DV", "CGIT_projects", "CGIT_amount", "CGIT_type"])


XCOLS = ["CGIT_projects", "CGIT_amount", "CGIT_type"]


# ---- manual FE (LSDV) helpers for small-cluster robustness (wild bootstrap, leave-one-out) ----
def _fe_fit(pdf, xcols, y):
    countries = sorted(pdf["country"].unique())
    D = np.column_stack([(pdf["country"].values == c).astype(float) for c in countries])
    XD = np.column_stack([pdf[xcols].values.astype(float), D])
    beta, *_ = np.linalg.lstsq(XD, y, rcond=None)
    return beta, XD, y - XD @ beta


def _cr_cov(pdf, XD, resid):
    XtXi = np.linalg.pinv(XD.T @ XD)
    meat = np.zeros((XD.shape[1], XD.shape[1]))
    cs = pdf["country"].values
    for c in np.unique(cs):
        idx = cs == c
        s = XD[idx].T @ resid[idx]
        meat += np.outer(s, s)
    N, K = XD.shape; G = len(np.unique(cs))
    return ((G / (G - 1)) * ((N - 1) / (N - K))) * (XtXi @ meat @ XtXi), G


def _analytic_p(pdf, xcols, j):
    y = pdf["DV"].values.astype(float)
    beta, XD, resid = _fe_fit(pdf, xcols, y)
    V, G = _cr_cov(pdf, XD, resid)
    se = np.sqrt(V[j, j]); t = beta[j] / se
    return beta[j], se, 2 * stats.t.sf(abs(t), G - 1)


def wcb_p(pdf, xcols, j, B=1999, seed=0):
    rng = np.random.default_rng(seed)
    y = pdf["DV"].values.astype(float)
    beta, XD, resid = _fe_fit(pdf, xcols, y)
    V, _ = _cr_cov(pdf, XD, resid)
    t_obs = beta[j] / np.sqrt(V[j, j])
    xr = [c for i, c in enumerate(xcols) if i != j]
    beta_r, XDr, resid_r = _fe_fit(pdf, xr, y)
    yhat_r = XDr @ beta_r
    cs = pdf["country"].values; uniq = np.unique(cs); cnt = 0
    for _ in range(B):
        w = {c: (1.0 if rng.random() < 0.5 else -1.0) for c in uniq}
        wv = np.array([w[c] for c in cs])
        bb, XDb, rb = _fe_fit(pdf, xcols, yhat_r + wv * resid_r)
        Vb, _ = _cr_cov(pdf, XDb, rb)
        if abs(bb[j] / np.sqrt(Vb[j, j])) >= abs(t_obs): cnt += 1
    return (cnt + 1) / (B + 1)


def hinges_on(pdf, xcols, j):
    var = xcols[j]
    n_ident = int((pdf[var] != 0).sum())
    txt = f"{n_ident} identifying row(s)"
    for c in sorted(pdf["country"].unique()):
        sub = pdf[pdf["country"] != c]
        if sub[var].nunique() <= 1:
            return txt + f"; drop {c} → no variation left"
        _, _, p = _analytic_p(sub, xcols, j)
        if p >= 0.05:
            return txt + f"; drop {c} → p={p:.2f}"
    return txt


def stars(p):
    return "***" if p < 0.001 else "**" if p < 0.01 else "*" if p < 0.05 else ""


def hausman(fe, re, names):
    try:
        b_fe = fe.params[names]; b_re = re.params[names]
        vd = fe.cov.loc[names, names].values - re.cov.loc[names, names].values
        d = (b_fe - b_re).values
        stat = float(d @ np.linalg.pinv(vd) @ d)
        k = len(names); p = float(stats.chi2.sf(stat, k))
        return stat, k, p
    except Exception as e:
        return None, None, None


def run_model(dv, dv_years, label):
    pdf = build_panel(dv, dv_years)
    df = pdf.set_index(["country", "year"])
    X = df[["CGIT_projects", "CGIT_amount", "CGIT_type"]]
    y = df["DV"]
    res = {"label": label, "N": int(len(df)), "G": int(pdf["country"].nunique()),
           "T": len(dv_years), "pdf": pdf}
    # FE (within) + RE, cluster-robust by country
    fe = PanelOLS(y, X, entity_effects=True).fit(cov_type="clustered", cluster_entity=True)
    Xc = X.assign(const=1.0)[["const", "CGIT_projects", "CGIT_amount", "CGIT_type"]]
    re = RandomEffects(y, Xc).fit(cov_type="clustered", cluster_entity=True)
    # Hausman on unadjusted covariances (classic form)
    fe_u = PanelOLS(y, X, entity_effects=True).fit()
    re_u = RandomEffects(y, Xc).fit()
    hs, hk, hp = hausman(fe_u, re_u, ["CGIT_projects", "CGIT_amount", "CGIT_type"])
    res.update(fe=fe, re=re, hausman=(hs, hk, hp))
    return res


VARROWS = [("CGIT_projects", "CGIT_projectsᵢ,ₜ₋₁"),
           ("CGIT_amount", "CGIT_amountᵢ,ₜ₋₁"),
           ("CGIT_type", "CGIT_typeᵢ,ₜ₋₁")]


def fstat_str(lm):
    """Overall model test H0: all slopes = 0 -- F for FE, Wald χ² for RE -- from the fitted result."""
    try:
        f = lm.f_statistic
        df = int(f.df)
        dfd = getattr(f, "df_denom", None)
        if dfd is not None and np.isfinite(dfd):
            return f"F({df},{int(dfd)})={f.stat:.2f}, p={f.pval:.3f}"
        return f"χ²({df})={f.stat:.2f}, p={f.pval:.3f}"
    except Exception:
        return "n/a"


def add_table(doc, res):
    doc.add_heading(res["label"], 2)
    fe, re = res["fe"], res["re"]
    t = doc.add_table(rows=1, cols=5); t.style = "Table Grid"
    hdr = ["Variable", "FE coefficient", "FE S.E.", "RE coefficient", "RE S.E."]
    for i, x in enumerate(hdr): t.rows[0].cells[i].text = x
    def cell(row, name, disp):
        c = t.add_row().cells; c[0].text = disp
        if name in fe.params.index:
            c[1].text = f"{fe.params[name]:.4f}{stars(fe.pvalues[name])}"; c[2].text = f"{fe.std_errors[name]:.4f}"
        else:
            c[1].text = "(absorbed)"; c[2].text = ""
        c[3].text = f"{re.params[name]:.4f}{stars(re.pvalues[name])}"; c[4].text = f"{re.std_errors[name]:.4f}"
    for name, disp in VARROWS:
        cell(t, name, disp)
    # constant (RE only; FE absorbs it)
    c = t.add_row().cells; c[0].text = "constant"
    c[1].text = "(absorbed)"; c[2].text = ""
    c[3].text = f"{re.params['const']:.4f}{stars(re.pvalues['const'])}"; c[4].text = f"{re.std_errors['const']:.4f}"
    # diagnostics
    def diag(lbl, fev, rev):
        c = t.add_row().cells; c[0].text = lbl; c[1].text = fev; c[3].text = rev
    diag("N (observations)", str(res["N"]), str(res["N"]))
    diag("N (groups)", str(res["G"]), str(res["G"]))
    diag("R² (within)", f"{fe.rsquared_within:.4f}", f"{re.rsquared_within:.4f}")
    diag("R² (overall)", f"{fe.rsquared_overall:.4f}", f"{re.rsquared_overall:.4f}")
    diag("F-statistic / Wald χ²", fstat_str(fe), fstat_str(re))
    hs, hk, hp = res["hausman"]
    hstr = f"χ²({hk})={hs:.2f}, p={hp:.3f}" if hs is not None else "degenerate (short panel)"
    doc.add_paragraph(f"Hausman test (FE vs RE): {hstr}. "
                      f"Panel: {res['G']} countries × T={res['T']}, N={res['N']}. "
                      f"Cluster-robust SEs by country. Significance: *** p<0.001, ** p<0.01, * p<0.05.").italic = True
    for row in t.rows:
        for cl in row.cells:
            for pa in cl.paragraphs:
                for rn in pa.runs: rn.font.size = Pt(8)


def add_robustness(doc, models):
    cols = []  # (short_label, var, res) for each coefficient significant in the main FE model
    for res in models:
        short = res["label"].split("—")[0].strip()
        for var in XCOLS:
            if var in res["fe"].pvalues.index and res["fe"].pvalues[var] < 0.05:
                cols.append((short, var, res))
    doc.add_heading("Robustness of significant coefficients (small-cluster inference)", 1)
    if not cols:
        doc.add_paragraph("No coefficient was significant at p<0.05 in the main models.")
        return
    t = doc.add_table(rows=1, cols=1 + len(cols)); t.style = "Table Grid"
    for i, (short, var, res) in enumerate(cols):
        t.rows[0].cells[i + 1].text = f"{short} · {var}"

    def row(name, vals):
        c = t.add_row().cells; c[0].text = name
        for i, v in enumerate(vals): c[i + 1].text = v

    naive, wcb, verdict, hinge = [], [], [], []
    for short, var, res in cols:
        pdf = res["pdf"]; j = XCOLS.index(var)
        naive.append(f"{res['fe'].pvalues[var]:.4f}")
        p = wcb_p(pdf, XCOLS, j)
        wcb.append(f"{p:.3f}")
        verdict.append("does not survive" if p >= 0.05 else "survives")
        hinge.append(hinges_on(pdf, XCOLS, j))
    row("naive cluster-robust p", naive)
    row("wild cluster bootstrap p", wcb)
    row("verdict (survives p<0.05?)", verdict)
    row("what it hinges on", hinge)
    doc.add_paragraph("Wild cluster bootstrap: 1,999 Rademacher draws, restricted null, cluster = country "
                      "(10 clusters) — the standard remedy for <40 clusters. 'What it hinges on' is "
                      "leave-one-country-out: the single country whose removal makes the coefficient "
                      "non-significant or unidentified.").italic = True
    for r in t.rows:
        for cl in r.cells:
            for pa in cl.paragraphs:
                for rn in pa.runs: rn.font.size = Pt(8)


def main():
    print("computing DVs (China share per country-year)...")
    cdn = nedko_cdn_china()
    ip = openintel_ip_china()
    print("running Model 1 (CDN, 2024-2026)...")
    m1 = run_model(cdn, [2024, 2025, 2026], "Model 1 — Chinese CDN market concentration (CDN_China), Nedko")
    print("running Model 2 (IP, 2025-2026)...")
    m2 = run_model(ip, [2025, 2026], "Model 2 — Chinese IP-infrastructure market concentration (IP_China), OpenINTEL")

    doc = Document()
    doc.add_heading("RQ3 — Panel regression: lagged DSR (CGIT) and Chinese infrastructure concentration", 0)
    doc.add_paragraph("EXPLORATORY / underpowered: short panels (Model 1 T=3, N=30; Model 2 T=2, N=20); "
                      "only ~5-6 countries have any DSR activity. Any statistically significant coefficient "
                      "should be interpreted with strong caution given the small N and sparse predictors.")
    for res in (m1, m2):
        add_table(doc, res)
        # console summary of significance
        fe, re = res["fe"], res["re"]
        sig = []
        for name in ["CGIT_projects", "CGIT_amount", "CGIT_type"]:
            if name in fe.pvalues.index and fe.pvalues[name] < 0.05: sig.append(f"FE {name} p={fe.pvalues[name]:.3f}")
            if re.pvalues[name] < 0.05: sig.append(f"RE {name} p={re.pvalues[name]:.3f}")
        print(f"  {res['label'][:40]}: significant (p<0.05) -> {sig if sig else 'NONE'}")

    print("computing robustness (wild cluster bootstrap + leave-one-out)...")
    add_robustness(doc, [m1, m2])

    out = os.path.join(BASE, "rq3_regression.docx")
    doc.save(out)
    print(f"\nDONE -> {out}")


if __name__ == "__main__":
    main()
