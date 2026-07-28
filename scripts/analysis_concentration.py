"""
Market-concentration analysis for the SEA CDN / IP-infrastructure study.

Produces, for each of the three sources (Nedko CDN, ISOC CDN, OpenINTEL IP infrastructure):
  (1) Gini by provider      (2) Gini by jurisdiction
  (3) HHI  by provider      (4) HHI  by jurisdiction
= 12 line graphs (one line per SEA country, x = time), each with a cross-tab share table,
all written into concentration_outputs/concentration_analysis.docx (+ the 12 PNGs).

MARKET-SHARE UNIT (agreed):
  Nedko     : 1 top-1000 domain = 1 row/month -> share by distinct domain per provider.
  OpenINTEL : distinct (website, provider) per country-month -- collapses the day x IP
              repetition (a website re-observed ~31 days x several IPs is still one website);
              a multi-homed website counts once for EACH provider it uses.
  ISOC      : website-count-weighted per provider-month, from the RAW lccrawresults (the
              formatted file dropped the counts). Multiple dates in a month -> latest date.

BASE: concentration among SERVED sites -- exclude Nedko 'unknown'/'None detected',
      OpenINTEL 'ASN ...' markers, ISOC 'summary'/'other'.

METRICS: Gini (standard, of the share distribution); HHI = sum(share^2) * 10000 (0-10000).
Jurisdiction metrics use the fixed 3-category market {China, USA, Other}.

>>> CONFIRM THIS ONE PARAMETER before running: whether ISOC weight is local, external or both.
"""
import csv, os, sys, math, datetime
from collections import defaultdict, Counter
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from docx import Document
from docx.shared import Inches, Pt

csv.field_size_limit(sys.maxsize)

BASE = "/Users/newlivehung/Desktop/11. Pulse Research Fellowship/06.07.2026_Folder"
RAW_ISOC = "/Users/newlivehung/Desktop/11. Pulse Research Fellowship/26.05.2026_Folder/lccrawresults.csv"
CCS = ["BN", "ID", "KH", "LA", "MM", "MY", "PH", "SG", "TH", "VN"]
OUTDIR = os.path.join(BASE, "concentration_outputs")
os.makedirs(OUTDIR, exist_ok=True)

# >>> CONFIRM: ISOC website weight -- "local", "external", or "local+external" <<<
ISOC_WEIGHT = "local+external"

TOP_PROVIDERS_IN_TABLE = 12   # OpenINTEL has 721 providers -> show top-N + "Other" in tables
DPI = 300                     # figure resolution (raise further for even crisper enlargement)
JUR3 = ["China", "USA", "Other"]

ISOC_JUR = {  # host_cdn -> 3-category jurisdiction (from provider_legal_names mapping)
    "cloudflare": "USA", "akamai": "USA", "amazon": "USA", "cloudfront": "USA", "fastly": "USA",
    "microsoft": "USA", "facebook": "USA", "netflix": "USA", "edgio": "USA", "wordpress": "USA",
    "alibaba": "China", "bunnycdn": "Other", "cdn77": "Other", "gcore": "Other",
}


def to_3cat(v):
    v = (v or "").strip()
    if v in ("China", "CN"): return "China"
    if v in ("US", "USA"): return "USA"
    if v == "" or v.upper() == "NULL": return ""
    return "Other"


# ---------- metrics ----------
def gini(weights):
    xs = sorted(weights)
    n = len(xs); s = sum(xs)
    if n == 0 or s <= 0: return float("nan")
    if n == 1: return 0.0
    cum = sum((i + 1) * x for i, x in enumerate(xs))
    return (2 * cum) / (n * s) - (n + 1) / n


def hhi(weights):
    s = sum(weights)
    if s <= 0: return float("nan")
    return sum((w / s) ** 2 for w in weights) * 10000


# self-test of the math (synthetic, touches no research data)
assert abs(gini([1, 1, 1, 1])) < 1e-9
assert abs(gini([1, 0, 0, 0]) - 0.75) < 1e-9
assert abs(hhi([1, 1]) - 5000) < 1e-9
assert abs(hhi([1, 0, 0, 0]) - 10000) < 1e-9


def period_to_date(p):
    return datetime.date(int(p[:4]), int(p[4:6]), 1)


# ---------- share builders: return prov_w[(cc,period)]=Counter, jur_w[(cc,period)]=Counter ----------
def build_nedko():
    prov = defaultdict(Counter); jur = defaultdict(Counter)
    for cc in CCS:
        path = os.path.join(BASE, "nedko_sea_output_mapped", f"nedko_{cc.lower()}_202401_202605.csv")
        with open(path, newline="", encoding="utf-8") as f:
            for row in csv.DictReader(f):
                p = row["cdn_provider"]
                if p in ("unknown", "None detected"): continue
                key = (cc, row["measurement_period"])
                prov[key][p] += 1
                c3 = to_3cat(row["cdn_jurisdiction"])
                if c3: jur[key][c3] += 1
    return prov, jur


def build_openintel():
    prov_sets = defaultdict(lambda: defaultdict(set)); jur_sets = defaultdict(lambda: defaultdict(set))
    for cc in CCS:
        path = os.path.join(BASE, "OpenINTEL_sea_output_caida_mapped", f"openintel_ip_infrastructure_extended_{cc}.csv")
        with open(path, newline="", encoding="utf-8") as f:
            r = csv.reader(f)
            h = next(r)
            i_mp = h.index("measurement_period"); i_web = h.index("website")
            i_prov = h.index("ip_infrastructure_provider"); i_jur = h.index("infrastructure_jurisdiction")
            for row in r:
                p = row[i_prov]
                if p.startswith("ASN "): continue
                key = (cc, row[i_mp]); w = row[i_web]
                prov_sets[key][p].add(w)
                c3 = to_3cat(row[i_jur])
                if c3: jur_sets[key][c3].add(w)
    prov = defaultdict(Counter); jur = defaultdict(Counter)
    for key, d in prov_sets.items():
        for p, s in d.items(): prov[key][p] = len(s)
    for key, d in jur_sets.items():
        for c, s in d.items(): jur[key][c] = len(s)
    return prov, jur


def build_isoc():
    tmp = {}   # (cc,month,host_cdn) -> (date_key, weight) keeping the latest date in the month
    with open(RAW_ISOC, newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            cc = (row.get("test_cc") or "").strip(); cdn = (row.get("host_cdn") or "").strip()
            if cc not in CCS or cdn in ("summary", "other") or not cdn: continue
            dd, mm, yyyy = row["date"].split("/")
            month = yyyy + mm.zfill(2); dkey = (yyyy, mm.zfill(2), dd.zfill(2))
            l = int(row["local_websites"] or 0); e = int(row["external_websites"] or 0)
            w = l if ISOC_WEIGHT == "local" else e if ISOC_WEIGHT == "external" else l + e
            k = (cc, month, cdn)
            if k not in tmp or dkey > tmp[k][0]:
                tmp[k] = (dkey, w)
    prov = defaultdict(Counter); jur = defaultdict(Counter)
    for (cc, month, cdn), (dkey, w) in tmp.items():
        key = (cc, month)
        prov[key][cdn] = w
        jur[key][ISOC_JUR.get(cdn, "Other")] += w
    return prov, jur


# ---------- per (cc,period) metric series ----------
def metric_series(weight_dict, dim):
    """dim='provider' -> over present providers; dim='jurisdiction' -> fixed 3-cat vector."""
    g = {}; h = {}
    for key, cnt in weight_dict.items():
        if dim == "jurisdiction":
            w = [cnt.get(c, 0) for c in JUR3]          # fixed 3-category market (keep zeros)
        else:
            w = [x for x in cnt.values() if x > 0]     # present providers only
        g[key] = gini(w); h[key] = hhi(w)
    return g, h


def yearly_shares(weight_dict):
    """cc -> key -> year -> mean monthly share (%), averaged over ALL months in that
    (country, year) — a key absent in a month contributes 0 — so each year's shares
    sum to ~100%."""
    monthly = {}                              # (cc,period) -> {key: share}
    cc_year_months = defaultdict(set)         # (cc,year)   -> {period}
    keys_in = defaultdict(set)                # (cc,year)   -> {key}
    for (cc, period), cnt in weight_dict.items():
        tot = sum(cnt.values())
        if tot <= 0: continue
        sh = {k: v / tot for k, v in cnt.items()}
        monthly[(cc, period)] = sh
        y = period[:4]
        cc_year_months[(cc, y)].add(period)
        keys_in[(cc, y)].update(sh.keys())
    out = defaultdict(lambda: defaultdict(dict))
    for (cc, y), months in cc_year_months.items():
        n = len(months)
        for k in keys_in[(cc, y)]:
            s = sum(monthly[(cc, p)].get(k, 0.0) for p in months)
            out[cc][k][y] = 100 * s / n
    return out


def top_keys(weight_dict, n):
    tot = Counter()
    for cnt in weight_dict.values():
        tot.update(cnt)
    return [k for k, _ in tot.most_common(n)]


def collapse_to_top(ysh, keep):
    keep = set(keep)
    out = defaultdict(lambda: defaultdict(lambda: defaultdict(float)))
    for cc, kd in ysh.items():
        for k, yd in kd.items():
            tgt = k if k in keep else "Other"
            for y, v in yd.items():
                out[cc][tgt][y] += v
    return out


# ---------- plotting ----------
def plot_metric(series, title, ylabel, fname):
    fig, ax = plt.subplots(figsize=(9, 4.4))
    for cc in CCS:
        pts = sorted((period_to_date(p), v) for (c, p), v in series.items()
                     if c == cc and not math.isnan(v))
        if pts:
            xs, ys = zip(*pts)
            ax.plot(xs, ys, marker="o", ms=2.5, lw=1.2, label=cc)
    ax.set_title(title, fontsize=12); ax.set_ylabel(ylabel); ax.grid(alpha=.25)
    ax.legend(ncol=5, fontsize=7, loc="best")
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m"))
    fig.autofmt_xdate(); fig.tight_layout()
    fig.savefig(fname, dpi=DPI); plt.close(fig)


# ---------- docx cross-tab ----------
def add_crosstab(doc, ysh, years, key_order, key_label):
    t = doc.add_table(rows=1, cols=2 + len(years))
    t.style = "Table Grid"   # plain bordered table (no shading / accent colours)
    hdr = t.rows[0].cells
    hdr[0].text = "Country"; hdr[1].text = key_label
    for i, y in enumerate(years): hdr[2 + i].text = y
    for cc in CCS:
        first = True
        for k in key_order:
            c = t.add_row().cells
            c[0].text = cc if first else ""
            c[1].text = k
            for i, y in enumerate(years):
                v = ysh.get(cc, {}).get(k, {}).get(y)
                c[2 + i].text = f"{v:.1f}" if v is not None else "–"
            first = False
    for row in t.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(7.5)


def main():
    print("building share tables (this streams the OpenINTEL files; ~1-2 min)...")
    sources = {}
    print("  Nedko...");     sources["Nedko (CDN)"] = build_nedko()
    print("  ISOC...");      sources["ISOC (CDN)"] = build_isoc()
    print("  OpenINTEL..."); sources["OpenINTEL (IP infrastructure)"] = build_openintel()

    doc = Document()
    doc.add_heading("SEA Internet-infrastructure market concentration", 0)
    doc.add_paragraph(f"Gini & HHI by provider and by jurisdiction, per country over time. "
                      f"Market-share unit: distinct website per provider (Nedko/OpenINTEL); "
                      f"website-count-weighted (ISOC, weight = {ISOC_WEIGHT}). "
                      f"HHI on the 0–10,000 scale. Jurisdiction = China / USA / Other.")

    plan = [("Gini", "provider"), ("Gini", "jurisdiction"), ("HHI", "provider"), ("HHI", "jurisdiction")]
    for metric, dim in plan:
        doc.add_heading(f"{metric} — by {dim}", 1)
        for sname, (prov_w, jur_w) in sources.items():
            wd = prov_w if dim == "provider" else jur_w
            g, h = metric_series(wd, dim)
            series = g if metric == "Gini" else h
            ylabel = "Gini coefficient" if metric == "Gini" else "HHI (0–10,000)"
            fname = os.path.join(OUTDIR, f"{metric}_{dim}_{sname.split()[0]}.png")
            plot_metric(series, f"{metric} — {dim} — {sname}", ylabel, fname)

            # share cross-tab
            ysh = yearly_shares(wd)
            years = sorted({y for cc in ysh.values() for k in cc.values() for y in k})
            if dim == "jurisdiction":
                korder = JUR3
            elif sname.startswith("OpenINTEL"):
                keep = top_keys(wd, TOP_PROVIDERS_IN_TABLE)
                ysh = collapse_to_top(ysh, keep); korder = keep + ["Other"]
            else:
                korder = top_keys(wd, 999)

            doc.add_heading(sname, 2)
            doc.add_picture(fname, width=Inches(6.3))
            cap = doc.add_paragraph(f"Figure: {metric} of {dim} concentration, {sname}, per country over time.")
            cap.runs[0].italic = True
            doc.add_paragraph(f"Table: market share (%) by {dim} × year × country ({sname}). "
                              f"Cells = mean monthly share in that year.").runs[0].italic = True
            add_crosstab(doc, ysh, years, korder,
                         "Provider" if dim == "provider" else "Jurisdiction")
            print(f"  [ok] {metric} / {dim} / {sname}")

    out = os.path.join(OUTDIR, "concentration_analysis.docx")
    doc.save(out)
    print(f"\nDONE -> {out}\n(12 figures also in {OUTDIR})")


if __name__ == "__main__":
    main()
